import os
import json
from typing import Dict, List, Optional, Any
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
import requests
from io import BytesIO
import base64

class Exporter:
    """
    Advanced export functionality for articles
    Handles DOCX, HTML, and analytics export with professional formatting
    """
    
    def __init__(self):
        """Initialize exporter with styling configurations"""
        
        # HTML template configuration
        self.html_template = """
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    <meta name="description" content="{meta_description}">
    <meta name="keywords" content="{keywords}">
    
    <!-- Open Graph Meta Tags -->
    <meta property="og:title" content="{title}">
    <meta property="og:description" content="{meta_description}">
    <meta property="og:type" content="article">
    
    <!-- Schema.org JSON-LD -->
    <script type="application/ld+json">
    {schema_markup}
    </script>
    
    <style>
        {css_styles}
    </style>
</head>
<body>
    <article class="ai-article">
        <header class="article-header">
            <h1 class="article-title">{title}</h1>
            <div class="article-meta">
                <p class="meta-description">{meta_description}</p>
                <div class="article-info">
                    <span class="word-count">{word_count} words</span>
                    <span class="read-time">{read_time} min read</span>
                </div>
            </div>
        </header>
        
            {cover_image}

            <div class="article-content">
                {content_sections}
            </div>
        
        <footer class="article-footer">
            {cta_section}
        </footer>
    </article>
</body>
</html>
        """.strip()
        
        # CSS styles for HTML export
        self.css_styles = """
        * {
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }
        
        body {
            font-family: 'Georgia', 'Times New Roman', serif;
            line-height: 1.6;
            color: #2c3e50;
            background-color: #f8f9fa;
            padding: 0;
        }
        
        .ai-article {
            max-width: 800px;
            margin: 0 auto;
            background: white;
            box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
        }
        
        .article-header {
            background: white;
            color: #2c3e50;
            padding: 3rem 2rem 2rem;
            text-align: center;
            border-bottom: 1px solid #e9ecef;
        }
        
        .article-title {
            font-size: 2.5rem;
            font-weight: 700;
            margin-bottom: 1rem;
            line-height: 1.2;
        }
        
        .meta-description {
            font-size: 1.1rem;
            color: #6c757d;
            margin-bottom: 1.5rem;
            font-style: italic;
        }
        
        .article-info {
            display: flex;
            justify-content: center;
            gap: 2rem;
            font-size: 0.9rem;
        }
        
        .article-info span {
            background: rgba(102, 126, 234, 0.1);
            color: #667eea;
            padding: 0.3rem 0.8rem;
            border-radius: 15px;
            border: 1px solid rgba(102, 126, 234, 0.2);
        }
        
        .article-content {
            padding: 3rem 2rem;
        }
        
        .content-section {
            margin-bottom: 3rem;
        }
        
        .section-heading {
            color: #34495e;
            font-size: 1.8rem;
            font-weight: 600;
            margin-bottom: 1rem;
            padding-bottom: 0.5rem;
            border-bottom: 3px solid #667eea;
        }
        
        .section-content {
            font-size: 1.1rem;
            line-height: 1.8;
            margin-bottom: 1.5rem;
        }
        
        .section-content p {
            margin-bottom: 1rem;
        }
        
        .section-image {
            margin: 2rem 0;
            text-align: center;
        }
        
        .section-image img {
            max-width: 100%;
            height: auto;
            border-radius: 10px;
            box-shadow: 0 4px 8px rgba(0, 0, 0, 0.1);
        }
        
        .image-caption {
            font-style: italic;
            color: #7f8c8d;
            margin-top: 0.5rem;
            font-size: 0.9rem;
        }
        
        .article-footer {
            background: #f8f9fa;
            padding: 2rem;
            border-top: 1px solid #e9ecef;
        }
        
        .cta-section {
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            padding: 2rem;
            border-radius: 10px;
            margin-bottom: 2rem;
            text-align: center;
        }
        
        .cta-section h3 {
            margin-bottom: 1rem;
            font-size: 1.5rem;
        }
        
        .ai-signature {
            text-align: center;
            color: #6c757d;
        }
        
        .ai-signature em {
            background: linear-gradient(45deg, #667eea, #764ba2);
            -webkit-background-clip: text;
            -webkit-text-fill-color: transparent;
            font-weight: bold;
        }
        
        a {
            color: #667eea;
            text-decoration: none;
            border-bottom: 1px solid transparent;
            transition: border-color 0.3s ease;
        }
        
        a:hover {
            border-bottom-color: #667eea;
        }
        
        @media (max-width: 768px) {
            .article-title {
                font-size: 2rem;
            }
            
            .article-header,
            .article-content,
            .article-footer {
                padding: 1.5rem;
            }
            
            .article-info {
                flex-direction: column;
                gap: 0.5rem;
            }
        }
        """
    
    def generate_docx(
        self, 
        article_data: Dict[str, Any], 
        seo_data: Dict[str, Any], 
        images: Optional[Dict[Any, Dict[str, str]]] = None
    ) -> bytes:
        """
        Generate professionally formatted DOCX document
        
        Args:
            article_data: The article content data
            seo_data: SEO analysis and metadata
            images: Generated images data
            
        Returns:
            DOCX file as bytes
        """
        
        # Create document
        doc = Document()
        
        # Configure document styles
        self._setup_docx_styles(doc)
        
        # Add document header with metadata
        self._add_docx_header(doc, article_data, seo_data)
        
        # Add title
        title_para = doc.add_heading(article_data.get('title', 'Generated Article'), level=1)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add meta description
        if article_data.get('meta_description'):
            meta_para = doc.add_paragraph(article_data['meta_description'])
            meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            meta_format = meta_para.runs[0].font
            meta_format.italic = True
            meta_format.size = Pt(12)
            meta_format.color.rgb = RGBColor(108, 117, 125)
        
        # Add separator
        doc.add_paragraph('─' * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add cover image (only one image per article). No caption for cover.
        if images:
            cover = None
            # Support both 'cover' string key and numeric 0 key for backward compatibility
            if isinstance(images, dict):
                cover = images.get('cover') or images.get(0) or images.get('0')
            if cover:
                self._add_image_to_docx(doc, cover, add_caption=False)
        
        # Add article sections
        for i, section in enumerate(article_data.get('sections', [])):
            # Add section heading
            if section.get('heading'):
                heading_para = doc.add_heading(section['heading'], level=2)
                heading_format = heading_para.runs[0].font
                heading_format.color.rgb = RGBColor(52, 73, 94)
            
            # Add section content
            if section.get('content'):
                # Clean and format content
                clean_content = self._clean_content_for_docx(section['content'])
                content_para = doc.add_paragraph(clean_content)
                content_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            # (Per-section images removed — exporter now uses a single cover image only)
            
            # Add spacing between sections
            doc.add_paragraph()
        
        # Add CTA section
        if article_data.get('cta'):
            doc.add_page_break()
            cta_heading = doc.add_heading('Call to Action', level=2)
            cta_para = doc.add_paragraph(article_data['cta'])
            cta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Style CTA
            for run in cta_para.runs:
                run.font.bold = True
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(102, 126, 234)
        
        # Add footer with metadata
        self._add_docx_footer(doc, seo_data)
        
        # Save to BytesIO
        doc_buffer = BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        
        return doc_buffer.getvalue()
    
    def _setup_docx_styles(self, doc: Document):
        """Setup custom styles for the DOCX document"""
        
        styles = doc.styles
        
        # Create custom paragraph style
        try:
            para_style = styles.add_style('CustomParagraph', WD_STYLE_TYPE.PARAGRAPH)
            para_format = para_style.paragraph_format
            para_format.space_after = Pt(12)
            para_format.line_spacing = 1.15
            
            font = para_style.font
            font.name = 'Calibri'
            font.size = Pt(11)
        except:
            pass  # Style might already exist
    
    def _add_docx_header(self, doc: Document, article_data: Dict[str, Any], seo_data: Dict[str, Any]):
        """Add document header with metadata"""
        
        # Document properties
        core_props = doc.core_properties
        core_props.title = article_data.get('title', 'Generated Article')
        core_props.author = 'AI Article Generator'
        core_props.subject = article_data.get('meta_description', '')[:100]
        core_props.keywords = ', '.join(seo_data.get('target_keywords', []))
        core_props.created = datetime.now()
        core_props.modified = datetime.now()
        
        # Add metadata table
        metadata_table = doc.add_table(rows=5, cols=2)
        metadata_table.style = 'Light Grid Accent 1'
        
        metadata = [
            ('Word Count', str(seo_data.get('word_count', 0))),
            ('SEO Score', f"{seo_data.get('seo_score', 0)}/100"),
            ('Readability', seo_data.get('readability_analysis', {}).get('reading_level', 'Standard')),
            ('Keywords', ', '.join(seo_data.get('target_keywords', [])[:3])),
            ('Generated', datetime.now().strftime('%Y-%m-%d %H:%M'))
        ]
        
        for i, (key, value) in enumerate(metadata):
            metadata_table.cell(i, 0).text = key
            metadata_table.cell(i, 1).text = value
        
        doc.add_paragraph()
    
    def _clean_content_for_docx(self, content: str) -> str:
        """Clean HTML and format content for DOCX"""
        
        import re
        
        # Remove HTML tags
        clean_content = re.sub(r'<[^>]+>', '', content)
        
        # Convert HTML entities
        html_entities = {
            '&amp;': '&',
            '&lt;': '<',
            '&gt;': '>',
            '&quot;': '"',
            '&#39;': "'",
            '&nbsp;': ' '
        }
        
        for entity, char in html_entities.items():
            clean_content = clean_content.replace(entity, char)
        
        # Clean up whitespace
        clean_content = re.sub(r'\s+', ' ', clean_content).strip()
        
        return clean_content
    
    def _add_image_to_docx(self, doc: Document, image_data: Dict[str, str], add_caption: bool = True):
        """Add image to DOCX document"""
        try:
            image_url = image_data.get('url', '')
            caption = image_data.get('caption', 'AI Generated Image')

            # Download image if it's a URL
            if image_url.startswith('http'):
                response = requests.get(image_url, timeout=10)
                if response.status_code == 200:
                    image_stream = BytesIO(response.content)

                    # Add image
                    paragraph = doc.add_paragraph()
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                    run.add_picture(image_stream, width=Inches(5))

                    # Add caption only if requested
                    if add_caption:
                        caption_para = doc.add_paragraph(caption)
                        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        caption_para.runs[0].font.italic = True
                        caption_para.runs[0].font.size = Pt(10)
                        caption_para.runs[0].font.color.rgb = RGBColor(108, 117, 125)

        except Exception as e:
            print(f"Failed to add image to DOCX: {e}")
            # Add placeholder text instead
            placeholder_para = doc.add_paragraph(f"[Image: {image_data.get('caption', 'AI Generated Image')}]")
            placeholder_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            placeholder_para.runs[0].font.italic = True
    
    def _add_docx_footer(self, doc: Document, seo_data: Dict[str, Any]):
        """Add footer with generation info"""
        
        doc.add_page_break()
        
        # AI signature
        signature_para = doc.add_paragraph()
        signature_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        signature_run = signature_para.add_run("✨ Generated by AI Writer — Powered by OpenAI GPT-4 & DALL-E 3 🎨")
        signature_run.font.italic = True
        signature_run.font.size = Pt(10)
        signature_run.font.color.rgb = RGBColor(102, 126, 234)
        
    def generate_html(
        self, 
        article_data: Dict[str, Any], 
        seo_data: Dict[str, Any], 
        images: Optional[Dict[int, Dict[str, str]]] = None
    ) -> str:
        """
        Generate clean, SEO-optimized HTML
        
        Args:
            article_data: The article content data
            seo_data: SEO analysis and metadata
            images: Generated images data
            
        Returns:
            Complete HTML document as string
        """
        
        # Prepare template variables
        template_vars = {
            'title': article_data.get('title', 'Generated Article'),
            'meta_description': article_data.get('meta_description', ''),
            'keywords': ', '.join(seo_data.get('target_keywords', [])),
            'word_count': seo_data.get('word_count', 0),
            'read_time': self._calculate_read_time(seo_data.get('word_count', 0)),
            'css_styles': self.css_styles,
            'schema_markup': self._generate_schema_markup(article_data, seo_data),
            'content_sections': self._generate_html_sections(article_data, images),
            'cover_image': self._generate_cover_html(images),
            'cta_section': self._generate_html_cta(article_data.get('cta', ''))
        }
        
        # Fill template
        html_content = self.html_template.format(**template_vars)
        
        return html_content
    
    def _calculate_read_time(self, word_count: int) -> int:
        """Calculate estimated reading time"""
        # Average reading speed: 200-250 words per minute
        return max(1, round(word_count / 225))
    
    def _generate_schema_markup(self, article_data: Dict[str, Any], seo_data: Dict[str, Any]) -> str:
        """Generate JSON-LD schema markup"""
        
        schema = {
            "@context": "https://schema.org",
            "@type": "Article",
            "headline": article_data.get('title', ''),
            "description": article_data.get('meta_description', ''),
            "author": {
                "@type": "Organization",
                "name": "AI Article Generator"
            },
            "publisher": {
                "@type": "Organization",
                "name": "AI Article Generator"
            },
            "datePublished": datetime.now().isoformat(),
            "dateModified": datetime.now().isoformat(),
            "wordCount": seo_data.get('word_count', 0),
            "keywords": seo_data.get('target_keywords', [])
        }
        
        return json.dumps(schema, indent=2)
    
    def _generate_html_sections(
        self, 
        article_data: Dict[str, Any], 
        images: Optional[Dict[int, Dict[str, str]]] = None
    ) -> str:
        """Generate HTML for article sections"""
        
        # Check if the full article content has HTML links (indicates link integration)
        full_content = article_data.get('content', '')
        has_integrated_links = '<a href=' in full_content
        
        if has_integrated_links:
            # Use the full content with integrated links instead of sections
            print("🔗 Exporter: Using full content with integrated links for HTML generation")
            
            # Process the full content that contains integrated links with proper formatting
            processed_content = self._process_full_content_with_links(full_content)
            
            # Wrap in a single content section
            return f'<div class="content-section"><div class="section-content">{processed_content}</div></div>'
        
        else:
            # Use original section-by-section processing
            sections_html = []
            
            for i, section in enumerate(article_data.get('sections', [])):
                section_html = []
                
                # Section wrapper
                section_html.append('<div class="content-section">')
                
                # Heading
                if section.get('heading'):
                    section_html.append(f'<h2 class="section-heading">{section["heading"]}</h2>')
                
                # Content
                if section.get('content'):
                    # Process content (convert simple formatting)
                    content = self._process_html_content(section['content'])
                    section_html.append(f'<div class="section-content">{content}</div>')
                
                # (Per-section images removed — exporter uses single cover image only)
                
                section_html.append('</div>')
                sections_html.append('\n'.join(section_html))
            
            return '\n\n'.join(sections_html)

    def _generate_cover_html(self, images: Optional[Dict[Any, Dict[str, str]]]) -> str:
        """Generate HTML for the single cover image (no caption)."""

        if not images or not isinstance(images, dict):
            return ""

        cover = images.get('cover') or images.get(0) or images.get('0')
        if not cover:
            return ""

        image_url = cover.get('url', '')
        alt_text = cover.get('alt_text', cover.get('caption', 'Cover image'))

        if cover.get('is_placeholder'):
            image_url = f"https://via.placeholder.com/1200x600/667eea/ffffff?text={cover.get('caption','Cover').replace(' ', '+')}"

        image_html = f"""
        <div class="cover-image" style="text-align:center; margin: 1.5rem 0;">
            <img src="{image_url}" alt="{alt_text}" loading="lazy" style="max-width:100%; height:auto; border-radius: 8px; box-shadow: 0 6px 18px rgba(0,0,0,0.12);">
        </div>
        """

        return image_html
    
    def _process_html_content(self, content: str) -> str:
        """Process and clean content for HTML output, preserving HTML links"""
        
        # Clean up markdown formatting issues but preserve HTML links
        content = content.replace('**', '').strip()
        
        # If content already contains HTML tags (like links), handle it more carefully
        if '<a href=' in content or '<img' in content or '<strong>' in content:
            # Content already has HTML formatting, preserve it and just ensure proper paragraph structure
            
            # First, check if content is already properly structured with <p> tags
            if '<p>' in content and '</p>' in content:
                # Content already has paragraph structure, return as-is
                return content
            
            # Split by double newlines for natural paragraphs
            if '\n\n' in content:
                paragraphs = content.split('\n\n')
            else:
                # If no double newlines, treat the whole content as one paragraph
                # but preserve internal structure
                paragraphs = [content]
        else:
            # No HTML formatting, use original logic
            paragraphs = content.split('\n\n')
        
        processed_paragraphs = []
        
        for para in paragraphs:
            para = para.strip()
            if para:
                # Don't double-wrap if already has paragraph tags or is a heading
                if not para.startswith('<p') and not para.startswith('<h') and not para.startswith('<div'):
                    # Check if paragraph contains HTML tags that should not be wrapped
                    if '<li>' in para or '<ul>' in para or '<ol>' in para or '<blockquote>' in para:
                        # List or blockquote content, don't wrap in <p>
                        processed_paragraphs.append(para)
                    else:
                        # Regular paragraph content, wrap in <p> tags
                        para = f'<p>{para}</p>'
                        processed_paragraphs.append(para)
                else:
                    processed_paragraphs.append(para)
        
        result = '\n'.join(processed_paragraphs)
        
        # Debug logging to track content processing
        link_count_before = content.count('<a href=')
        link_count_after = result.count('<a href=')
        print(f"📄 HTML Content Processing: {link_count_before} links → {link_count_after} links")
        
        return result
    
    def _process_full_content_with_links(self, content: str) -> str:
        """Process full content with integrated links, handling markdown formatting properly"""
        import re
        
        # Clean up any markdown formatting but preserve HTML links
        content = content.replace('**', '').strip()
        
        # Split by markdown headings and process each section
        sections = re.split(r'^##\s+(.+)$', content, flags=re.MULTILINE)
        
        # First section is usually intro content before first heading
        processed_sections = []
        
        for i, section in enumerate(sections):
            section = section.strip()
            if not section:
                continue
                
            if i == 0:
                # First section - intro content
                if section:
                    paragraphs = section.split('\n\n')
                    for para in paragraphs:
                        para = para.strip()
                        if para:
                            processed_sections.append(f'<p>{para}</p>')
            elif i % 2 == 1:
                # Odd indices are headings
                processed_sections.append(f'<h2 class="section-heading">{section}</h2>')
            else:
                # Even indices are content after headings
                paragraphs = section.split('\n\n')
                for para in paragraphs:
                    para = para.strip()
                    if para:
                        processed_sections.append(f'<p>{para}</p>')
        
        result = '\n'.join(processed_sections)
        
        # Debug logging
        link_count = result.count('<a href=')
        heading_count = result.count('<h2')
        print(f"📄 Full Content Processing: {link_count} links preserved, {heading_count} headings formatted")
        
        return result
    
    def _generate_image_html(self, image_data: Dict[str, str]) -> str:
        """Generate HTML for image display"""
        
        image_url = image_data.get('url', '')
        caption = image_data.get('caption', '')
        alt_text = image_data.get('alt_text', caption)
        
        if image_data.get('is_placeholder'):
            # For placeholder images, use a more professional placeholder service
            image_url = f"https://via.placeholder.com/800x400/667eea/ffffff?text={caption.replace(' ', '+')}"
        
        image_html = f"""
        <div class="section-image">
            <img src="{image_url}" alt="{alt_text}" loading="lazy">
            <div class="image-caption">{caption}</div>
        </div>
        """
        
        return image_html
    
    def _generate_html_cta(self, cta_content: str) -> str:
        """Generate HTML for CTA section"""
        
        if not cta_content:
            return ""
        
        # Skip CTA if it contains unwanted promotional content
        if "Ready to Take Action?" in cta_content or "Generated by AI Writer" in cta_content:
            return ""
        
        cta_html = f"""
        <div class="cta-section">
            <p>{cta_content}</p>
        </div>
        """
        
        return cta_html
    
    def generate_analytics(
        self, 
        article_data: Dict[str, Any], 
        seo_data: Dict[str, Any]
    ) -> str:
        """
        Generate comprehensive analytics JSON
        
        Args:
            article_data: The article content data
            seo_data: SEO analysis and metadata
            
        Returns:
            Analytics data as JSON string
        """
        
        analytics = {
            'generation_metadata': {
                'timestamp': datetime.now().isoformat(),
                'generator': 'AI Article Generator v1.0',
                'article_id': self._generate_article_id(article_data),
                'language': article_data.get('language', 'English'),
                'tone': article_data.get('tone', 'Professional')
            },
            
            'content_metrics': {
                'title': article_data.get('title', ''),
                'total_word_count': seo_data.get('word_count', 0),
                'section_count': len(article_data.get('sections', [])),
                'estimated_read_time_minutes': self._calculate_read_time(seo_data.get('word_count', 0)),
                'section_word_distribution': [
                    {
                        'section_index': i,
                        'heading': section.get('heading', ''),
                        'word_count': section.get('word_count', len(section.get('content', '').split()))
                    }
                    for i, section in enumerate(article_data.get('sections', []))
                ]
            },
            
            'seo_metrics': {
                'overall_score': seo_data.get('seo_score', 0),
                'focus_keyword': seo_data.get('focus_keyword', ''),
                'target_keywords': seo_data.get('target_keywords', []),
                'keyword_density': seo_data.get('keyword_density', 0),
                'title_optimization': seo_data.get('title_analysis', {}),
                'meta_description_optimization': seo_data.get('meta_analysis', {}),
                'readability_score': seo_data.get('readability_score', 0),
                'recommendations_count': len(seo_data.get('recommendations', [])),
                'lsi_keywords_suggested': seo_data.get('lsi_suggestions', [])
            },
            
            'content_quality': {
                'readability_analysis': seo_data.get('readability_analysis', {}),
                'structure_analysis': seo_data.get('content_structure', {}),
                'link_analysis': {
                    'internal_links': len([l for l in seo_data.get('links', []) if l.get('type') == 'internal']),
                    'external_links': len([l for l in seo_data.get('links', []) if l.get('type') == 'external']),
                    'total_links': len(seo_data.get('links', []))
                }
            },
            
            'optimization_suggestions': {
                'high_priority': [r for r in seo_data.get('recommendations', [])[:3]],
                'medium_priority': [r for r in seo_data.get('recommendations', [])[3:6]],
                'low_priority': [r for r in seo_data.get('recommendations', [])[6:]]
            },
            
            'export_info': {
                'formats_available': ['HTML', 'DOCX', 'JSON'],
                'seo_ready': seo_data.get('seo_score', 0) >= 70,
                'publication_ready': self._assess_publication_readiness(article_data, seo_data)
            }
        }
        
        return json.dumps(analytics, indent=2, ensure_ascii=False)
    
    def _generate_article_id(self, article_data: Dict[str, Any]) -> str:
        """Generate unique article ID"""
        
        title = article_data.get('title', 'article')
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        title_hash = hash(title) % 10000
        
        return f"art_{timestamp}_{title_hash}"
    
    def _assess_publication_readiness(
        self, 
        article_data: Dict[str, Any], 
        seo_data: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Assess if article is ready for publication"""
        
        readiness_score = 0
        max_score = 100
        issues = []
        
        # Content completeness (30 points)
        if article_data.get('title'):
            readiness_score += 10
        else:
            issues.append("Missing title")
        
        if article_data.get('meta_description'):
            readiness_score += 10
        else:
            issues.append("Missing meta description")
        
        if len(article_data.get('sections', [])) >= 3:
            readiness_score += 10
        else:
            issues.append("Need at least 3 content sections")
        
        # SEO optimization (40 points)
        seo_score = seo_data.get('seo_score', 0)
        readiness_score += int(seo_score * 0.4)
        
        if seo_score < 60:
            issues.append("SEO score needs improvement")
        
        # Content quality (30 points)
        word_count = seo_data.get('word_count', 0)
        if word_count >= 300:
            readiness_score += 15
        elif word_count >= 200:
            readiness_score += 10
        else:
            issues.append("Article too short (needs 300+ words)")
        
        readability_score = seo_data.get('readability_score', 0)
        if readability_score >= 60:
            readiness_score += 15
        elif readability_score >= 40:
            readiness_score += 10
        else:
            issues.append("Content readability needs improvement")
        
        return {
            'score': min(max_score, readiness_score),
            'is_ready': readiness_score >= 80,
            'issues': issues,
            'recommendations': [
                "Review and address all identified issues",
                "Consider adding more relevant keywords",
                "Ensure proper internal linking",
                "Add compelling call-to-action"
            ] if issues else ["Article is ready for publication!"]
        }
    
    def export_all_formats(
        self, 
        article_data: Dict[str, Any], 
        seo_data: Dict[str, Any], 
        images: Optional[Dict[int, Dict[str, str]]] = None,
        output_dir: str = "exports"
    ) -> Dict[str, str]:
        """
        Export article in all available formats
        
        Returns:
            Dictionary mapping format names to file paths
        """
        
        # Ensure output directory exists
        os.makedirs(output_dir, exist_ok=True)
        
        # Generate base filename
        title_slug = self._slugify(article_data.get('title', 'article'))
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        base_filename = f"{title_slug}_{timestamp}"
        
        export_paths = {}
        
        try:
            # Export HTML
            html_content = self.generate_html(article_data, seo_data, images)
            html_path = os.path.join(output_dir, f"{base_filename}.html")
            with open(html_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
            export_paths['HTML'] = html_path
            
            # Export DOCX
            docx_content = self.generate_docx(article_data, seo_data, images)
            docx_path = os.path.join(output_dir, f"{base_filename}.docx")
            with open(docx_path, 'wb') as f:
                f.write(docx_content)
            export_paths['DOCX'] = docx_path
            
            # Export Analytics JSON
            analytics_content = self.generate_analytics(article_data, seo_data)
            analytics_path = os.path.join(output_dir, f"{base_filename}_analytics.json")
            with open(analytics_path, 'w', encoding='utf-8') as f:
                f.write(analytics_content)
            export_paths['Analytics'] = analytics_path
            
        except Exception as e:
            print(f"Export error: {e}")
        
        return export_paths
    
    def _slugify(self, text: str) -> str:
        """Convert text to URL-friendly slug"""
        
        import re
        
        slug = text.lower()
        slug = re.sub(r'[^\w\s-]', '', slug)
        slug = re.sub(r'\s+', '_', slug)
        slug = re.sub(r'-+', '-', slug)
        slug = slug.strip('-_')
        
        return slug[:30] if len(slug) > 30 else slug